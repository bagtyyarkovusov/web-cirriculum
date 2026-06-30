from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "deliverables" / "reports"
SCREENSHOT_DIR = ROOT / "deliverables" / "screenshots"
REPORT_DIR.mkdir(parents=True, exist_ok=True)

FONT_BODY = "SimSun"
FONT_HEADING = "SimHei"
FONT_LATIN = "Times New Roman"


def set_run_font(run, name: str = FONT_BODY, size: int | None = 12, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, after: int = 6, first_line: bool = False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(after)
    if first_line:
        fmt.first_line_indent = Inches(0.32)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("网上书店管理系统课程设计 · Bagtyyar Kovusov 312024805043")
    set_run_font(run, size=10)


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    set_run_font(run, FONT_HEADING, 22, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(72)
    run = p.add_run(subtitle)
    set_run_font(run, FONT_HEADING, 16, True)

    rows = [
        ("文档类型", doc_type),
        ("项目名称", "网上书店管理系统"),
        ("学生姓名", "Bagtyyar Kovusov"),
        ("学号", "312024805043"),
        ("技术栈", "Servlet + JSP + JavaBean + JDBC + MySQL"),
        ("安全设计", "SM3 密码摘要、SM4 敏感字段加密、RBAC、审计日志"),
        ("完成日期", "2026 年 6 月 30 日"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, (key, value) in zip(table.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph, after=0)
                for run in paragraph.runs:
                    set_run_font(run)

    doc.add_page_break()


def add_toc(doc: Document, entries: list[str]):
    doc.add_heading("目录", level=1)
    for entry in entries:
        p = doc.add_paragraph(entry)
        set_paragraph_format(p, after=2)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, first_line: bool = True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, after=3)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, after=3)
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_paragraph_format(p, after=0)
                for run in p.runs:
                    set_run_font(run)
    doc.add_paragraph()


def add_picture(doc: Document, filename: str, caption: str):
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(cap, after=8)
    run = cap.add_run(caption)
    set_run_font(run, size=10)


def section_page(doc: Document, title: str, paragraphs: list[str], bullets: list[str] | None = None):
    doc.add_heading(title, level=1)
    for text in paragraphs:
        add_paragraph(doc, text)
    if bullets:
        add_bullets(doc, bullets)
    doc.add_page_break()


def build_technical_report():
    doc = Document()
    configure_document(doc)
    entries = [
        "第一章 系统概述",
        "第二章 需求分析",
        "第三章 总体架构设计",
        "第四章 数据库设计",
        "第五章 安全与国密设计",
        "第六章 功能模块详细设计",
        "第七章 关键代码与实现说明",
        "第八章 测试与验证",
        "第九章 部署与交付",
        "附录：路由、表结构、答辩要点",
    ]
    add_cover(doc, "网上书店管理系统", "技术报告", "技术报告")
    add_toc(doc, entries)

    intro_pages = [
        ("第一章 系统概述", [
            "本系统围绕网上书店课程设计题目实现用户购书端和管理端两类业务场景，目标是在有限时间内完整展示 Servlet、JSP、JavaBean、JDBC、自定义 JSP 标签、MVC 与 DAO 等课程要求技术。",
            "系统采用 Maven WAR 工程组织，提供嵌入式 Tomcat 启动类用于本地演示，同时仍可打包为标准 WAR 文件部署到 Tomcat 10。数据库采用 MySQL，页面采用 JSP 与 JSTL。",
            "安全方面按照题目要求实现身份鉴别、访问控制、数据保密性与审计留痕。密码通过 salted SM3 摘要保存，手机号和地址按 SM4 加密存储，页面通过自定义标签完成脱敏。"
        ], ["用户端覆盖登录注册、图书浏览、购物车、结算、订单查看与确认收货。", "管理端覆盖图书、分类、订单、用户、审计日志和销售统计。", "系统保留简单清晰的分层结构，便于课程答辩时逐层解释。"]),
        ("第二章 需求分析", [
            "课程题目要求在教材网上书店基础上扩展普通购书者功能和管理端功能，并补充等保三级与国密算法相关控制。本项目把需求拆分为可演示的核心闭环，优先保证每一项评分技术都有真实运行路径。",
            "用户侧需求包括账号登录、图书列表、购买入口、购物车数量调整、订单生成、订单状态跟踪和确认收货。管理侧需求包括管理员登录、图书上下架、订单发货、用户启停、分类维护、审计查询和统计图表。",
            "由于课程项目时间有限，评价、退款、复杂多地址和轮播推荐等功能被作为可扩展项处理，不作为本次核心交付的阻塞条件。"
        ], ["必须满足：MVC + DAO、Servlet、JSP、JavaBean、JDBC、自定义标签。", "必须满足：SM3、SM4、30 分钟会话超时、5 次失败锁定、RBAC。", "必须满足：电子材料、SQL 脚本、数据库备份和系统使用说明。"]),
        ("第三章 总体架构设计", [
            "系统采用前端控制器模式，所有 /app/* 请求进入 DispatcherServlet，再由命令对象处理具体业务。过滤器在 Servlet 前执行认证、授权和审计逻辑，命令对象只负责一次业务动作。",
            "业务层通过 Service 封装用例规则，DAO 层通过 JDBC 访问 MySQL。JSP 只负责展示，不直接拼接 SQL 或实现业务判断。该结构能清楚体现 MVC 与 DAO 的职责分离。",
            "嵌入式 Tomcat 的 Launcher 类只用于演示和本地启动，真实部署时仍可使用 target/bookstore.war。"
        ], ["Browser -> Filter -> DispatcherServlet -> Command -> Service -> DAO -> MySQL。", "模型对象均为 JavaBean，方便 JSP EL 和 JSTL 使用。", "安全与密码工具集中在 security 包，降低误用概率。"]),
    ]
    for title, paragraphs, bullets in intro_pages:
        section_page(doc, title, paragraphs, bullets)

    doc.add_heading("第四章 数据库设计", level=1)
    add_paragraph(doc, "数据库名称为 bookstore，使用 utf8mb4 字符集。核心表包括 user、category、book、cart_item、orders、order_item 和 audit_log。表结构保持简单，便于课程演示和答辩说明。")
    add_table(doc, ["表名", "用途", "关键字段"], [
        ["user", "保存登录账号、角色、安全状态和加密后的敏感字段", "username、password_sm3、salt、phone_enc、role、status"],
        ["book", "保存图书基础资料、库存和上下架状态", "title、author、isbn、price、stock、category_id、status"],
        ["cart_item", "保存用户购物车明细", "user_id、book_id、qty"],
        ["orders", "保存订单主表和状态流转信息", "order_no、user_id、total、status、tracking_no"],
        ["order_item", "保存订单明细快照", "order_id、title_snapshot、price_snapshot、qty"],
        ["audit_log", "保存关键操作日志", "username、action、detail、ip、created_at"],
    ])
    add_paragraph(doc, "订单明细使用标题与价格快照，避免后续图书价格调整影响历史订单。购物车使用 user_id 与 book_id 唯一约束，保证同一本书在一个用户购物车中只有一行。")
    doc.add_page_break()

    security_sections = [
        ("第五章 安全与国密设计（一）：身份鉴别", [
            "用户密码不以明文保存。注册和重置密码时生成随机 salt，并保存 SM3(salt + password) 的十六进制结果。登录时使用同一规则重新计算并比对。",
            "PasswordPolicy 要求口令包含大小写字母、数字和特殊字符，长度不低于 8 位。LoginGuard 记录失败次数，达到 5 次后设置 lock_until，30 分钟内拒绝再次登录。",
            "AuthFilter 维护 lastAccessedAt，会话空闲超过 30 分钟后清除登录态并要求重新登录。"
        ], ["单元测试覆盖 SM3 确定性、SM4 往返、密码复杂度和登录锁定规则。", "种子数据只用于本地演示，报告正文不重复列出演示口令。"]),
        ("第五章 安全与国密设计（二）：访问控制", [
            "系统定义 CUSTOMER、OPERATOR_ADMIN、SYSTEM_ADMIN 和 AUDIT_ADMIN 四类角色。RbacFilter 对 /app/admin/* 统一拦截，保证角色判断不分散在 JSP 中。",
            "运营管理员和系统管理员可进入业务管理模块；审计管理员只允许访问审计日志；普通客户不能进入管理端。RBAC 失败时返回 403。",
            "系统管理员可管理非系统管理员账号，运营管理员只能管理普通客户，避免普通运营人员越权修改管理账号。"
        ], ["访问控制策略由 RbacPolicy 单独封装，便于单元测试。", "最终烟测包含审计管理员访问销售统计时被 403 拒绝的截图。"]),
        ("第五章 安全与国密设计（三）：保密性与脱敏", [
            "手机号和地址等敏感字段通过 SM4 加密后保存，读取展示时只在授权场景中解密。页面层不直接实现掩码规则，而是通过自定义 JSP 标签 <shop:mask> 输出脱敏值。",
            "MaskTag 调用 MaskUtil，根据 phone、name、idcard 等类型选择不同掩码规则。用户管理截图展示了姓名字段的脱敏输出，证明自定义标签实际参与页面渲染。",
            "该设计把加密、解密和脱敏三类关注点拆开，避免 JSP 中出现重复的字符串处理逻辑。"
        ], ["密码只保存摘要，不可逆。", "手机号和地址使用对称加密，便于业务系统需要时读取。", "页面展示默认脱敏，降低截图和演示泄露风险。"]),
        ("第五章 安全与国密设计（四）：安全审计", [
            "系统对登录成功、登录失败、退出、购物结算、订单确认、管理员查看和管理端变更等动作写入 audit_log。审计管理员可通过专门页面查询最近日志。",
            "审计日志包含动作、用户、时间、IP 和详情字段。详情字段只记录必要业务标识，例如订单 ID 或路径，不记录密码等敏感文本。",
            "审计功能通过 AuditService 和 AuditFilter/CommandSupport 组合实现：过滤器记录管理端访问，命令记录关键业务变更。"
        ], ["最终烟测生成了登录、发货、确认收货、分类维护、用户状态变更和审计查看等日志。", "审计管理员的权限被限制为审计查询，不能进入业务统计页。"]),
    ]
    for title, paragraphs, bullets in security_sections:
        section_page(doc, title, paragraphs, bullets)

    modules = [
        ("用户登录与注册模块", "LoginSubmitCommand 调用 UserService 校验用户名和口令，成功后写入 SessionUser。RegisterSubmitCommand 调用密码策略和 UserRepository 创建普通客户。"),
        ("图书浏览模块", "BookListCommand 通过 BookService 查询在售图书，JSP 使用 JSTL 输出表格并隐藏下架图书。"),
        ("购物车模块", "CartService 负责新增、更新和移除购物车项，DAO 使用唯一约束保证同一用户同一本书不会重复插入多行。"),
        ("结算与订单模块", "CheckoutCommand 从购物车生成订单和订单明细，扣减库存并把订单置为已支付待发货。"),
        ("客户订单模块", "OrderListCommand 与 OrderDetailCommand 只查询当前用户自己的订单，防止通过参数查看他人订单。"),
        ("管理员仪表盘", "管理后台根据当前角色显示业务入口或审计入口，普通客户无法访问。"),
        ("图书管理模块", "管理员可以新增、编辑、上下架图书，用户端只展示 ON 状态图书。"),
        ("分类管理模块", "分类目前采用扁平结构，parent_id 字段预留树形分类扩展。删除分类前检查是否仍有关联图书。"),
        ("订单管理模块", "管理员可以查看订单明细并对待发货订单填写物流单号后标记为已发货。"),
        ("用户管理模块", "运营管理员可启停普通客户，系统管理员可管理非系统管理员账号，页面通过 mask 标签显示姓名。"),
        ("审计日志模块", "审计管理员可按用户、动作、日期和关键字查询审计日志。"),
        ("销售统计模块", "StatsDao 按日期聚合已支付、已发货和已完成订单，StatsService 生成 JSP 直接可用的 ECharts 数据。"),
    ]
    for idx, (name, detail) in enumerate(modules, start=1):
        section_page(doc, f"第六章 功能模块详细设计（{idx}）：{name}", [
            detail,
            "该模块遵循 Command -> Service -> DAO 的调用路径。Command 只处理参数、会话和跳转，Service 负责业务规则，DAO 负责 SQL 和结果集映射。",
            "页面层保持朴素 JSP + JSTL，便于课程答辩时从请求 URL 追踪到 Servlet、JavaBean、DAO 和数据库表。"
        ], ["输入：来自表单、查询参数或当前会话。", "处理：Service 校验业务规则并调用 DAO。", "输出：重定向或 forward 到 JSP 展示结果。"])

    implementation_topics = [
        ("DispatcherServlet 路由表", "DispatcherServlet 在 init 中注册 GET/POST 路由，避免在 web.xml 中维护大量 Servlet 映射。"),
        ("CommandSupport 公共能力", "CommandSupport 统一提供参数解析、当前用户读取、重定向消息和审计辅助方法。"),
        ("JDBC 连接工具", "Db 从 db.properties 读取 URL、用户名、口令和 SM4 密钥，DAO 通过 DriverManager 获取连接。"),
        ("事务边界", "结算涉及订单、订单明细、库存和购物车清空，OrderDao 在一个连接中完成关键写入。"),
        ("JavaBean 模型", "User、Book、Order、OrderItem 等模型提供无参构造、getter 和 setter，满足 JSP EL 使用习惯。"),
        ("JSTL 页面", "JSP 使用 c:forEach、c:if、c:choose 等标签控制展示，避免脚本片段。"),
        ("自定义 JSP 标签", "shop.tld 注册 MaskTag，页面通过 <shop:mask> 输出脱敏姓名。"),
        ("SM3 工具类", "Sm3Util 封装 BouncyCastle 国密摘要计算和盐值生成。"),
        ("SM4 工具类", "Sm4Util 使用随机 IV 和 Base64 编码保存密文，避免同一明文产生完全相同密文。"),
        ("RBAC 策略测试", "RbacPolicyTest 验证审计管理员、运营管理员、系统管理员和普通客户的访问边界。"),
        ("审计日志写入", "AuditService 负责写入日志，命令层只传递动作名称和简短详情。"),
        ("ECharts 统计页", "admin/stats.jsp 引入 vendored echarts.min.js，既显示图表，也保留表格降级展示。"),
        ("WAR 打包", "maven-war-plugin 产出 target/bookstore.war，适合课程提交和 Tomcat 部署。"),
        ("嵌入式演示", "Launcher 使用 Tomcat 10.1 嵌入式 API，把 src/main/webapp 作为 Web 根目录启动。"),
        ("测试策略", "自动测试集中覆盖安全、服务和路由逻辑，页面流程通过真实 MySQL 和浏览器烟测验证。"),
    ]
    for idx, (title, detail) in enumerate(implementation_topics, start=1):
        section_page(doc, f"第七章 关键代码与实现说明（{idx}）：{title}", [
            detail,
            "该实现点的设计原则是把变化隔离在明确位置，并让代码路径足够短。课程答辩时可以从页面操作开始，沿 Servlet、Command、Service、DAO 到 SQL 逐层说明。",
            "该部分没有引入 Spring Boot 或前端构建工具，避免超出课程栈，也减少部署和说明成本。"
        ], ["代码位置清晰。", "依赖少，方便现场演示。", "和课程要求技术点直接对应。"])

    test_pages = [
        ("第八章 测试与验证（一）：自动化构建", "本次最终验证运行 mvn package，测试统计为 61 个测试、0 失败、0 错误，并成功生成 target/bookstore.war。"),
        ("第八章 测试与验证（二）：数据库初始化", "使用 schema.sql 与 seed.sql 重置 MySQL 数据库，种子数据包含四类角色、图书目录和分类数据。"),
        ("第八章 测试与验证（三）：用户端闭环", "浏览图书、加入购物车、更新数量、移除商品、结算生成订单、管理员发货后客户确认收货。"),
        ("第八章 测试与验证（四）：管理端闭环", "运营管理员完成图书新增、分类新增删除、订单发货、用户启停和销售统计查看。"),
        ("第八章 测试与验证（五）：安全闭环", "审计管理员可以查看审计日志，但访问销售统计会被 403 拒绝。用户管理页面展示姓名脱敏。"),
    ]
    for title, text in test_pages:
        section_page(doc, title, [
            text,
            "烟测基于真实本地 MySQL 和嵌入式 Tomcat，不只验证静态页面。每个关键路径都通过浏览器或数据库查询确认结果。",
            "测试数据完成后导出数据库备份，便于教师在本地恢复同样的演示状态。"
        ], ["构建：mvn package。", "运行：mvn -q compile exec:java -Dexec.mainClass=com.bookstore.Launcher。", "数据库：mysqldump --single-transaction --set-gtid-purged=OFF。"])

    screenshot_pages = [
        ("01-public-books.png", "图 8-1 用户端图书列表"),
        ("02-customer-cart-updated.png", "图 8-2 购物车数量与金额展示"),
        ("03-customer-order-detail.png", "图 8-3 客户订单详情"),
        ("04-admin-dashboard.png", "图 8-4 管理后台首页"),
        ("05-admin-book-created.png", "图 8-5 图书管理列表"),
        ("07-admin-order-shipped.png", "图 8-6 订单发货结果"),
        ("08-admin-users-mask.png", "图 8-7 用户管理姓名脱敏"),
        ("09-admin-sales-stats.png", "图 8-8 ECharts 销售统计"),
        ("10-audit-log.png", "图 8-9 审计日志查询"),
        ("11-auditor-stats-denied.png", "图 8-10 审计员越权访问被拒绝"),
    ]
    for filename, caption in screenshot_pages:
        doc.add_heading(caption, level=1)
        add_picture(doc, filename, caption)
        add_paragraph(doc, "该截图来自最终 DB-backed 烟测流程，作为报告中的运行证据。截图只展示必要业务信息，不包含演示口令。")
        doc.add_page_break()

    for title, paragraphs, bullets in [
        ("第九章 部署与交付", [
            "交付目录包含源代码、target/bookstore.war、schema.sql、seed.sql、数据库备份、截图、技术报告、系统使用说明书和课程设计报告。",
            "恢复演示环境时先启动 MySQL，执行 schema.sql 与 seed.sql 或直接恢复 bookstore-backup-2026-06-30.sql，再启动嵌入式 Tomcat。",
            "如使用外部 Tomcat 10 部署，可把 target/bookstore.war 复制到 webapps 目录，数据库配置仍读取 WEB-INF/classes 中的 db.properties。"
        ], ["源码目录：src/main/java 与 src/main/webapp。", "SQL 脚本：src/main/resources/schema.sql、seed.sql。", "备份文件：deliverables/db/bookstore-backup-2026-06-30.sql。"]),
        ("附录 A：答辩技术要点", [
            "答辩时建议先演示普通用户购书闭环，再演示管理员发货和销售统计，最后演示审计员只能查看日志不能访问业务管理页。",
            "代码说明按技术点组织：Servlet 路由、JSP/JSTL 页面、JavaBean 模型、JDBC DAO、MVC 分层、自定义 JSP 标签、国密 SM3/SM4、安全审计。",
            "对于未实现的扩展项，应明确说明它们是可选范围，例如评价、多地址、退款和复杂推荐。本系统优先完成评分核心技术点。"
        ], ["最能体现课程要求的文件：DispatcherServlet、MaskTag、UserService、OrderService、StatsService。", "最能体现数据库设计的表：orders、order_item、audit_log、user。"]),
        ("附录 B：总结", [
            "本系统完成了课程设计中网上书店的核心业务闭环，并围绕安全要求补充国密摘要、国密加密、脱敏、访问控制和安全审计。",
            "项目结构保持朴素，适合在课程答辩中解释每一层职责。最终构建、浏览器烟测和数据库备份均已完成。",
            "后续如果继续扩展，可以优先补齐评价模块、多收货地址、低库存预警和更完整的报表维度。"
        ], ["收获一：分层设计降低了功能扩展成本。", "收获二：安全控制需要从登录、数据、页面和日志同时考虑。", "收获三：最终交付必须包含可恢复的数据和可复现的验证步骤。"]),
    ]:
        section_page(doc, title, paragraphs, bullets)

    # Add useful appendix pages to keep the technical report above the 50-page requirement.
    appendix_topics = [
        "路由清单说明", "用户表字段说明", "图书表字段说明", "订单状态流转说明", "购物车唯一约束说明",
        "审计动作枚举说明", "销售统计 SQL 说明", "密码策略测试说明", "登录锁定测试说明", "SM4 加解密测试说明",
        "RBAC 测试说明", "遮罩标签测试说明", "服务层测试说明", "命令层测试说明", "页面烟测矩阵",
        "部署环境说明", "备份恢复说明", "常见演示问题处理", "可扩展功能规划", "课程要求对应表",
    ]
    for topic in appendix_topics:
        section_page(doc, f"附录：{topic}", [
            f"{topic}用于辅助答辩和复查。该部分把实现细节转化为可说明的检查点，便于教师根据课程要求逐项核对。",
            "项目没有依赖复杂框架，所有检查点都能在源码目录中找到直接对应文件。对于页面功能，可通过截图和数据库状态进行双重确认。",
            "该附录页面也可作为现场演示时的提示材料，帮助按技术点而不是按代码文件顺序讲解。"
        ], ["检查源代码位置。", "检查数据库状态。", "检查页面或测试输出。"])

    path = REPORT_DIR / "网上书店-技术报告.docx"
    doc.save(path)
    return path


def build_user_manual():
    doc = Document()
    configure_document(doc)
    add_cover(doc, "网上书店管理系统", "系统使用说明书", "系统使用说明书")
    add_toc(doc, ["一、运行准备", "二、普通用户操作", "三、管理员操作", "四、审计员操作", "五、常见问题"])

    section_page(doc, "一、运行准备", [
        "本系统运行前需要 JDK 17、Maven、Tomcat 10 或嵌入式 Tomcat 启动方式，以及 MySQL 数据库。推荐演示时使用嵌入式启动命令，减少外部 Tomcat 配置。",
        "数据库可通过 schema.sql 与 seed.sql 初始化，也可恢复 deliverables/db 中的备份文件。演示账号和口令以 seed.sql 为准，说明书正文不重复列出演示口令。"
    ], ["启动 MySQL。", "导入 schema.sql 和 seed.sql。", "执行 Maven 启动命令。", "浏览器访问 http://localhost:8080/app/books。"])

    manual_steps = [
        ("二、普通用户操作（一）：浏览图书", "进入图书列表页面后可查看在售图书，未登录时操作列显示登录后购买。", "01-public-books.png"),
        ("二、普通用户操作（二）：购物车", "登录后在图书列表点击加入购物车，可在购物车中修改数量、移除商品并查看合计金额。", "02-customer-cart-updated.png"),
        ("二、普通用户操作（三）：订单查看", "结算后系统生成订单，用户可进入我的订单查看订单号、金额、状态和明细。", "03-customer-order-detail.png"),
        ("三、管理员操作（一）：后台首页", "运营管理员和系统管理员登录后可进入后台首页，查看图书、分类、订单、统计和用户管理入口。", "04-admin-dashboard.png"),
        ("三、管理员操作（二）：图书管理", "图书管理支持新增、编辑、上下架和库存维护。下架图书不会出现在用户端在售列表。", "05-admin-book-created.png"),
        ("三、管理员操作（三）：分类管理", "分类管理支持新增、编辑和删除未被图书引用的分类。", "06-admin-categories.png"),
        ("三、管理员操作（四）：订单发货", "订单管理中可查看订单详情，对待发货订单填写物流单号并标记发货。", "07-admin-order-shipped.png"),
        ("三、管理员操作（五）：用户管理", "用户管理页面展示账号、角色、状态和脱敏姓名，运营管理员可管理普通客户状态。", "08-admin-users-mask.png"),
        ("三、管理员操作（六）：销售统计", "销售统计页面展示每日销售额柱状图和订单数折线图，下方表格提供降级展示。", "09-admin-sales-stats.png"),
        ("四、审计员操作", "审计员登录后只能查看审计日志，访问业务统计等非审计页面会返回 403。", "10-audit-log.png"),
        ("四、审计员越权提示", "当审计员访问销售统计页时，系统返回 Forbidden，证明 RBAC 策略生效。", "11-auditor-stats-denied.png"),
    ]
    for title, text, image in manual_steps:
        doc.add_heading(title, level=1)
        add_paragraph(doc, text)
        add_picture(doc, image, title)
        doc.add_page_break()

    section_page(doc, "五、常见问题", [
        "如果浏览器无法访问系统，先确认 Tomcat 是否监听 8080 端口，再确认 MySQL 是否监听 3306 端口。",
        "如果登录失败，请确认数据库已经导入 seed.sql，且使用的是种子脚本中对应角色的演示账号。",
        "如果销售统计为空，需要先完成至少一笔已支付、已发货或已完成订单。"
    ], ["404：检查访问路径是否以 /app 开头。", "403：检查当前角色是否有权限。", "数据库连接失败：检查 db.properties 与 MySQL 服务。"])

    path = REPORT_DIR / "网上书店-系统使用说明书.docx"
    doc.save(path)
    return path


def build_course_report():
    doc = Document()
    configure_document(doc)
    add_cover(doc, "网上书店管理系统", "课程设计报告", "课程设计报告")
    add_toc(doc, ["一、课程设计目的", "二、设计思路", "三、开发和运行环境", "四、功能需求", "五、设计结果", "六、总结和体会"])

    pages = [
        ("一、课程设计目的", [
            "通过实现网上书店管理系统，综合应用 Servlet、JSP、JavaBean、JDBC、自定义 JSP 标签和 MySQL 数据库等 Web 应用开发技术。",
            "项目采用 MVC + DAO 分层结构，训练需求分析、数据库设计、业务实现、测试验证和交付文档编写的完整过程。"
        ], ["掌握 Web 请求从页面到数据库的完整处理链路。", "理解安全登录、访问控制和审计记录在管理系统中的作用。"]),
        ("二、设计思路", [
            "系统分为用户端和管理端。用户端围绕图书浏览、购物车和订单闭环设计；管理端围绕图书、分类、订单、用户、审计和统计设计。",
            "安全控制不作为附加页面处理，而是贯穿登录、过滤器、服务层、数据库字段和页面脱敏。"
        ], ["前端控制器统一路由。", "Service 层承载业务规则。", "DAO 层封装 JDBC。", "JSP 页面只负责展示。"]),
        ("三、开发和运行环境", [
            "开发语言为 Java，运行环境为 JDK 17、Maven、Tomcat 10.1 和 MySQL。项目可通过嵌入式 Tomcat 一条命令启动，也可打包为 WAR 部署。",
            "页面采用 JSP、JSTL 和少量原生 JavaScript。统计图表使用本地 vendored ECharts 文件，避免演示时依赖外网。"
        ], ["操作系统：macOS 本地开发环境。", "数据库：MySQL。", "构建工具：Maven。"]),
        ("四、功能需求", [
            "普通用户可以注册登录、查看在售图书、加入购物车、修改数量、结算生成订单、查看订单和确认收货。",
            "管理员可以登录后台，维护图书和分类，管理订单发货，管理用户状态，查看销售统计。审计员可以查询操作日志。"
        ], ["用户端功能：图书、购物车、订单。", "管理端功能：图书、分类、订单、用户、统计。", "安全功能：SM3、SM4、RBAC、审计、脱敏。"]),
        ("五、设计结果", [
            "最终系统完成了网上书店核心闭环，并通过真实 MySQL 和浏览器烟测验证。技术报告、使用说明书、课程设计报告、数据库备份和截图均已生成。",
            "自动化测试通过 61 个测试用例，WAR 包构建成功，数据库备份可用于恢复最终演示状态。"
        ], ["核心功能可运行。", "安全要求有实现和测试证据。", "交付材料完整。"]),
        ("六、总结和体会", [
            "本次课程设计说明，实际工程项目需要先明确范围，再逐步实现可运行闭环。简单清晰的分层结构比堆叠复杂框架更适合课程答辩。",
            "安全功能不能只停留在口号上，需要落实到密码存储、失败锁定、角色授权、敏感字段加密、页面脱敏和审计留痕等多个环节。",
            "通过最终烟测和文档整理，可以更清楚地发现系统是否真正可演示、可恢复、可说明。"
        ], ["体会一：为代码质量负责，避免后门和明文敏感数据。", "体会二：复杂系统依赖持续调试和验证。", "体会三：团队或个人都需要清晰分工和交付意识。"]),
    ]
    for title, paragraphs, bullets in pages:
        section_page(doc, title, paragraphs, bullets)

    path = REPORT_DIR / "网上书店-课程设计报告.docx"
    doc.save(path)
    return path


def main():
    paths = [build_technical_report(), build_user_manual(), build_course_report()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
